#!/usr/bin/env python3
"""
meeting_notes.py
----------------
Generate or supplement meeting notes DOCX from Zoom recordings:
  - Zoom closed captions (.txt)
  - Zoom chat transcript (.txt)
  - Audio recordings (.m4a / .mp3 / .wav) — transcribed via mlx-whisper

GENERATE MODE (no existing notes):
    python3 meeting_notes.py [options]

SUPPLEMENT MODE (existing notes DOCX as first argument):
    python3 meeting_notes.py <notes.docx> [options]

Options:
    --captions FILE [FILE] One or more closed caption text files
    --chat FILE [FILE]     One or more chat log text files
    --audio FILE [FILE] One or more audio files (transcribed with mlx-whisper)
    --output FILE       Output DOCX path
                          generate: default meeting_notes.docx
                          supplement: default <input>_supplemented.docx
    --title TITLE       Meeting title — generate mode only (default: auto-detected or "Meeting Notes")
    --date DATE         Meeting date — generate mode only (default: today)
    --model MODEL       mlx-whisper model (default: mlx-community/whisper-large-v3-turbo)
    --no-ai             Skip Claude AI; write raw transcripts into document

Requirements:
    pip install python-docx mlx-whisper

Examples:
    # Generate from scratch
    python3 meeting_notes.py \\
        --captions closed_caption.txt \\
        --chat chat.txt \\
        --audio audio.m4a \\
        --output meeting_notes.docx

    # Supplement an existing document
    python3 meeting_notes.py EVNPC_agenda.docx \\
        --captions closed_caption.txt \\
        --chat chat.txt \\
        --audio audio1.m4a audio2.m4a
"""

import argparse
import subprocess
import sys
import json
import re
from datetime import date
from pathlib import Path


# ── shared helpers ────────────────────────────────────────────────────────────

def read_text_file(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def transcribe_audio(audio_path, model="mlx-community/whisper-large-v3-turbo"):
    """Transcribe an audio file using mlx-whisper. Returns timestamped transcript string."""
    try:
        import mlx_whisper
    except ImportError:
        print("  mlx_whisper not installed. Run: pip install mlx-whisper", file=sys.stderr)
        return None

    print(f"  Transcribing {Path(audio_path).name} ...")
    result = mlx_whisper.transcribe(
        str(audio_path),
        path_or_hf_repo=model,
        word_timestamps=False,
    )
    lines = []
    for seg in result.get("segments", []):
        start = seg.get("start", 0)
        text = seg.get("text", "").strip()
        if text:
            lines.append(f"[{start:.1f}s] {text}")
    return "\n".join(lines)


def gather_sources(args):
    """Read captions, chat, and audio files into a dict of {label: text}."""
    sources = {}

    if args.captions:
        combined = []
        for f in args.captions:
            print(f"Reading captions: {f}")
            combined.append(f"--- {Path(f).name} ---\n{read_text_file(f)}")
        sources["Closed Captions"] = "\n\n".join(combined)

    if args.chat:
        combined = []
        for f in args.chat:
            print(f"Reading chat: {f}")
            combined.append(f"--- {Path(f).name} ---\n{read_text_file(f)}")
        sources["Chat Log"] = "\n\n".join(combined)

    if args.audio:
        combined = []
        for audio_file in args.audio:
            apath = Path(audio_file)
            cache_path = apath.with_suffix(".transcript.txt")
            alt_cache = apath.parent / (apath.stem + "_transcript.txt")
            if cache_path.exists():
                print(f"Using cached transcript: {cache_path.name}")
                transcript = read_text_file(cache_path)
            elif alt_cache.exists():
                print(f"Using cached transcript: {alt_cache.name}")
                transcript = read_text_file(alt_cache)
            else:
                transcript = transcribe_audio(str(apath), model=args.model)
                if transcript:
                    cache_path.write_text(transcript, encoding="utf-8")
                    print(f"  Saved transcript cache: {cache_path.name}")
            if transcript:
                combined.append(f"--- {apath.name} ---\n{transcript}")
        if combined:
            sources["Audio Transcript"] = "\n\n".join(combined)

    return sources


def _sources_block(sources, max_chars=400000):
    block = ""
    for label, content in sources.items():
        if content:
            block += f"\n\n=== {label.upper()} ===\n{content[:max_chars]}"
    return block


def _run_claude(prompt):
    """Run the Claude CLI with the given prompt. Returns stdout or None on error."""
    import shutil
    if not shutil.which("claude"):
        print("  'claude' CLI not found. Ensure Claude Code is installed and on PATH.", file=sys.stderr)
        return None
    print("  Calling Claude CLI...")
    result = subprocess.run(
        ["claude", "-p"],
        input=prompt,
        capture_output=True,
        text=True,
        timeout=300,
    )
    if result.returncode != 0:
        print(f"  Claude CLI error: {result.stderr.strip()}", file=sys.stderr)
        return None
    return result.stdout.strip()


# ── generate mode ─────────────────────────────────────────────────────────────

GENERATE_SYSTEM_PROMPT = """You are an expert meeting secretary. You will receive raw content from a Zoom meeting — closed captions, chat logs, and/or audio transcripts — and must produce structured meeting notes.

Output a single JSON object with the following keys:

{
  "title": "Meeting title (infer from content or use 'Meeting Notes')",
  "date": "Date if found in content, otherwise null",
  "attendees": ["list of names/identifiers mentioned"],
  "summary": "2-4 sentence executive summary of the meeting",
  "agenda_items": [
    {
      "heading": "Topic or agenda item title",
      "discussion": "Key points discussed (2-6 sentences)",
      "decisions": ["Decision 1", "Decision 2"],
      "action_items": [
        {"owner": "Name or 'TBD'", "action": "What they will do"}
      ]
    }
  ],
  "other_notes": ["Any important points not covered above"],
  "next_meeting": "Date/time of next meeting if mentioned, else null"
}

Rules:
- Be concise but complete. Capture specific names, numbers, decisions.
- Group related discussion into logical agenda items even if no formal agenda existed.
- Extract action items with owners wherever possible.
- Output JSON only, no preamble or trailing text.
"""


MAP_SYSTEM_PROMPT = """You are an expert scientific meeting secretary. You will receive ONE PORTION (a time-ordered segment) of a longer meeting's raw content — closed captions, chat, and/or audio transcript. This is not the whole meeting; earlier/later segments are handled separately.

Your job: EXHAUSTIVELY extract every distinct topic discussed in THIS portion. Do NOT compress the meeting into a few high-level bullets — that loses detail. Err on the side of MORE agenda items with specific content rather than fewer summary items.

Output a single JSON object:

{
  "attendees": ["names/identifiers that speak or are mentioned in this portion"],
  "agenda_items": [
    {
      "heading": "Specific topic / proposal / agenda item title",
      "discussion": "What was actually said — capture ALL specifics: proposal codes, grades, telescope names, frequencies, numbers, names, disagreements, rationale. Multiple sentences are fine.",
      "decisions": ["Concrete decisions made"],
      "action_items": [{"owner": "Name or 'TBD'", "action": "What they will do"}]
    }
  ],
  "next_meeting": "Date/time of next meeting if mentioned here, else null"
}

Rules:
- One agenda item per distinct topic/proposal. If ten proposals are discussed, produce ten items.
- Preserve every proposal code, grade, telescope, and number exactly as spoken.
- Do not editorialize or drop content for brevity. Completeness over concision.
- Output JSON only, no preamble or trailing text.
"""

REDUCE_SUMMARY_PROMPT = """You are an expert meeting secretary. Below are the agenda-item headings extracted from a full meeting. Write a JSON object with a title and an executive summary.

{
  "title": "Concise, specific meeting title",
  "summary": "3-6 sentence executive summary of the whole meeting"
}

Output JSON only.
"""


def _parse_json_object(response):
    """Extract the first top-level JSON object from a Claude response."""
    if not response:
        return None
    m = re.search(r'\{.*\}', response, re.DOTALL)
    if not m:
        print("  Warning: No JSON object found in Claude response.", file=sys.stderr)
        return None
    try:
        return json.loads(m.group())
    except json.JSONDecodeError as e:
        print(f"  Warning: Could not parse Claude JSON: {e}", file=sys.stderr)
        print(f"  Raw response (first 500 chars): {response[:500]}", file=sys.stderr)
        return None


def _chunk_text(text, chunk_chars=45000, overlap=2000):
    """Split text into overlapping, time-ordered chunks."""
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        chunks.append(text[i:i + chunk_chars])
        if i + chunk_chars >= n:
            break
        i += chunk_chars - overlap
    return chunks


# Above this many source chars, summarize in chunks so no single pass has to
# compress the whole meeting (which silently drops per-topic detail).
CHUNK_THRESHOLD = 55000


def call_claude_generate(sources):
    """Call Claude to produce meeting notes JSON from raw sources.

    For long meetings, use map-reduce chunking so detail isn't lost to a single
    over-compressed summarization pass.
    """
    block = _sources_block(sources)
    if len(block) <= CHUNK_THRESHOLD:
        prompt = f"{GENERATE_SYSTEM_PROMPT}\n\nSOURCE CONTENT:{block}\n\nPlease produce the meeting notes JSON object."
        return _parse_json_object(_run_claude(prompt))

    return _call_claude_generate_chunked(block)


def _call_claude_generate_chunked(block):
    """Map-reduce summarization for long meetings."""
    chunks = _chunk_text(block)
    print(f"  Long meeting ({len(block):,} chars): summarizing in {len(chunks)} chunks...")

    all_items = []
    attendees = []
    next_meeting = None

    for idx, chunk in enumerate(chunks, 1):
        print(f"  Chunk {idx}/{len(chunks)} ...")
        prompt = (
            f"{MAP_SYSTEM_PROMPT}\n\n"
            f"MEETING PORTION {idx} of {len(chunks)}:\n{chunk}\n\n"
            f"Please produce the JSON object for THIS portion."
        )
        data = _parse_json_object(_run_claude(prompt))
        if not data:
            print(f"  Warning: chunk {idx} produced no usable output; skipping.", file=sys.stderr)
            continue
        all_items.extend(data.get("agenda_items", []) or [])
        for a in data.get("attendees", []) or []:
            if a and a not in attendees:
                attendees.append(a)
        next_meeting = next_meeting or data.get("next_meeting")

    if not all_items:
        print("  Warning: no agenda items extracted from any chunk.", file=sys.stderr)
        return None

    # Reduce: derive a title + overall summary from the collected headings.
    headings = "\n".join(f"- {it.get('heading', '')}" for it in all_items if it.get("heading"))
    summ = _parse_json_object(
        _run_claude(f"{REDUCE_SUMMARY_PROMPT}\n\nAGENDA HEADINGS:\n{headings}\n\nProduce the JSON.")
    ) or {}

    return {
        "title": summ.get("title"),
        "date": None,
        "attendees": attendees,
        "summary": summ.get("summary", ""),
        "agenda_items": all_items,
        "other_notes": [],
        "next_meeting": next_meeting,
    }


def build_docx_from_notes(data, meeting_date, meeting_title):
    """Build a python-docx Document from structured notes dict."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = data.get("title") or meeting_title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    detected_date = data.get("date") or meeting_date
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_para.add_run(f"Date: {detected_date}")
    run.font.size = Pt(11)
    run.font.bold = True

    attendees = data.get("attendees", [])
    if attendees:
        att_para = doc.add_paragraph()
        att_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        att_para.add_run("Attendees: " + ", ".join(attendees)).font.size = Pt(10)

    doc.add_paragraph()

    summary = data.get("summary", "")
    if summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(summary)
        doc.add_paragraph()

    agenda_items = data.get("agenda_items", [])
    if agenda_items:
        doc.add_heading("Discussion", level=1)
        for item in agenda_items:
            heading = item.get("heading", "")
            if heading:
                doc.add_heading(heading, level=2)

            discussion = item.get("discussion", "")
            if discussion:
                doc.add_paragraph(discussion)

            decisions = item.get("decisions", [])
            if decisions:
                doc.add_paragraph().add_run("Decisions:").bold = True
                for d in decisions:
                    doc.add_paragraph(d, style="List Bullet")

            action_items = item.get("action_items", [])
            if action_items:
                doc.add_paragraph().add_run("Action Items:").bold = True
                for a in action_items:
                    owner = a.get("owner", "TBD")
                    action = a.get("action", "")
                    doc.add_paragraph(f"[{owner}] {action}", style="List Bullet")

            doc.add_paragraph()

    other = data.get("other_notes", [])
    if other:
        doc.add_heading("Other Notes", level=1)
        for note in other:
            doc.add_paragraph(note, style="List Bullet")
        doc.add_paragraph()

    next_meeting = data.get("next_meeting")
    if next_meeting:
        doc.add_heading("Next Meeting", level=1)
        doc.add_paragraph(next_meeting)

    return doc


def build_raw_docx(sources, meeting_date, meeting_title):
    """Fallback: dump raw transcripts into a document."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title_para = doc.add_heading(meeting_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Date: {meeting_date}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for label, content in sources.items():
        if not content:
            continue
        doc.add_heading(label, level=1)
        for line in content.split("\n"):
            if line.strip():
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    return doc


# ── supplement mode ───────────────────────────────────────────────────────────

SUPPLEMENT_SYSTEM_PROMPT = """You are an expert scientific meeting secretary helping to supplement EVN (European VLBI Network) Programme Committee meeting notes.

You will receive:
1. The existing meeting notes (from a DOCX agenda/notes document)
2. Supplementary sources: closed captions, chat log, and/or audio transcripts

Your task: For each section of the meeting notes, identify important details, decisions, discussions, or context that appear in the supplementary sources but are NOT already captured in the notes. Produce a structured JSON list of insertions.

Rules:
- Only add genuinely new information not already in the notes
- Be concise but complete — include specific names, numbers, decisions, action items
- Preserve technical terms (telescope names, proposal codes, grades, frequencies)
- Note the source type (captions/chat/audio)
- Output JSON only, no preamble

Output format:
[
  {
    "anchor": "exact substring of the notes paragraph to insert after",
    "note": "The supplementary text to insert (1-3 sentences, starting with [Source])",
    "source": "captions|chat|audio"
  },
  ...
]
"""


def call_claude_supplement(notes_text, sources):
    """Call Claude to produce insertion suggestions for an existing notes document."""
    prompt = (
        f"{SUPPLEMENT_SYSTEM_PROMPT}\n\n"
        f"MEETING NOTES:\n{notes_text[:200000]}\n\n"
        f"SUPPLEMENTARY SOURCES:{_sources_block(sources)}\n\n"
        f"Please produce the JSON insertion list as described."
    )
    response = _run_claude(prompt)
    if not response:
        return []
    json_match = re.search(r'\[.*\]', response, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError as e:
            print(f"  Warning: Could not parse Claude JSON: {e}", file=sys.stderr)
            print(f"  Raw response: {response[:500]}", file=sys.stderr)
    return []


def make_note_para(note_text, source="captions"):
    """Create a formatted supplementary note paragraph element."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    color_map = {"captions": "1F4E79", "chat": "7B2C8C", "audio": "375623"}
    icon_map  = {"captions": "📝",      "chat": "💬",      "audio": "🎙"}
    color = color_map.get(source, "1F4E79")
    icon  = icon_map.get(source, "📝")

    p    = OxmlElement('w:p')
    r    = OxmlElement('w:r')
    rpr  = OxmlElement('w:rPr')

    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    rpr.append(color_el)
    rpr.append(OxmlElement('w:i'))

    for tag, val in [('w:sz', '18'), ('w:szCs', '18')]:
        el = OxmlElement(tag)
        el.set(qn('w:val'), val)
        rpr.append(el)

    r.append(rpr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = f"{icon} {note_text}"
    r.append(t)
    p.append(r)
    return p


def apply_insertions(doc, insertions):
    """Apply a list of {anchor, note, source} dicts to the document."""
    inserted = 0
    for item in insertions:
        anchor = item.get("anchor", "")
        note   = item.get("note", "")
        source = item.get("source", "captions")
        if not anchor or not note:
            continue
        target = next((p for p in doc.paragraphs if anchor in p.text), None)
        if target is None:
            print(f"  Could not find anchor: '{anchor[:60]}'", file=sys.stderr)
            continue
        target._element.addnext(make_note_para(note, source))
        inserted += 1
    return inserted


def append_full_transcripts(doc, sources):
    """Append full transcripts as a new section at end of document."""
    from docx.shared import Pt

    doc.add_page_break()
    doc.add_heading("Supplementary Source Transcripts", level=1)
    for label, content in sources.items():
        if not content:
            continue
        doc.add_heading(label, level=2)
        for chunk in content.split('\n'):
            if chunk.strip():
                p = doc.add_paragraph(chunk)
                p.runs[0].font.size = Pt(8)


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("notes", nargs="?",
                        help="Existing meeting notes DOCX to supplement (omit to generate from scratch)")
    parser.add_argument("--captions", nargs="+", help="Closed captions text file(s)")
    parser.add_argument("--chat",     nargs="+", help="Chat log text file(s)")
    parser.add_argument("--audio", nargs="+", help="Audio files to transcribe")
    parser.add_argument("--output",   help="Output DOCX path")
    parser.add_argument("--title",    default="Meeting Notes",
                        help="Meeting title — generate mode only")
    parser.add_argument("--date",     default=str(date.today()),
                        help="Meeting date — generate mode only")
    parser.add_argument("--model",    default="mlx-community/whisper-large-v3-turbo",
                        help="Whisper model for audio transcription")
    parser.add_argument("--no-ai",    action="store_true",
                        help="Skip Claude AI; write raw transcripts into document")
    args = parser.parse_args()

    sources = gather_sources(args)

    # ── supplement mode ──
    if args.notes:
        notes_path = Path(args.notes)
        if not notes_path.exists():
            print(f"Error: notes file not found: {notes_path}", file=sys.stderr)
            sys.exit(1)

        output_path = Path(args.output) if args.output \
            else notes_path.with_stem(notes_path.stem + "_supplemented")

        if not sources:
            print("Warning: no supplementary sources provided. Nothing to add.", file=sys.stderr)
            sys.exit(0)

        from docx import Document
        print(f"\nLoading notes document: {notes_path.name}")
        doc = Document(str(notes_path))
        notes_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        if args.no_ai:
            print("Appending full transcripts to document (no-AI mode)...")
            append_full_transcripts(doc, sources)
        else:
            insertions = call_claude_supplement(notes_text, sources)
            if insertions:
                print(f"  Got {len(insertions)} insertion suggestions from Claude.")
                n = apply_insertions(doc, insertions)
                print(f"  Successfully inserted {n} notes into document.")
            else:
                print("  No insertions returned. Falling back to appending transcripts.")
                append_full_transcripts(doc, sources)

    # ── generate mode ──
    else:
        output_path = Path(args.output) if args.output else Path("meeting_notes.docx")

        if not sources:
            print("Error: no sources provided. Use --captions, --chat, and/or --audio.", file=sys.stderr)
            sys.exit(1)

        if args.no_ai:
            print("Building raw transcript document (no-AI mode)...")
            doc = build_raw_docx(sources, args.date, args.title)
        else:
            data = call_claude_generate(sources)
            if data:
                print("  Building structured meeting notes document...")
                doc = build_docx_from_notes(data, args.date, args.title)
            else:
                print("  Falling back to raw transcript document.")
                doc = build_raw_docx(sources, args.date, args.title)

    doc.save(str(output_path))
    print(f"\n✓ Saved to: {output_path}")


if __name__ == "__main__":
    main()
